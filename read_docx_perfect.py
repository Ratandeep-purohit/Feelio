from docx import Document
import sys

doc = Document(sys.argv[1])
with open(sys.argv[2], 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        f.write(p.text + '\n')
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            f.write("\t".join(row_data) + '\n')
